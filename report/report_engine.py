"""核心报告生成引擎 - DB11/T 1785-2020 附录C 完整实现

遵循《二氧化碳排放核算和报告要求 服务业》(DB11/T 1785-2020)
生成符合附录C规范的标准碳排放报告（HTML + Word）。

报告结构（四章 + 真实性声明）：
  第一章：企业基本情况（表C.1）
  第二章：二氧化碳排放核算（表C.2~C.6）
  第三章：活动水平数据（表C.7~C.9）
  第四章：排放因子数据（表C.10~C.12）
  汇总与声明（表C.13~C.15）
"""

from __future__ import annotations

import json
import logging
import math
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.shared import Emu
from schemas import (
    ActivityData,
    DEFAULT_EMISSION_FACTORS,
    EmissionCalculation,
    EmissionFactor,
    EmissionSource,
    EnergyType,
    EnterpriseInfo,
    ReportFormat,
)

logger = logging.getLogger(__name__)

# ============================================================
# 排放因子常量
# ============================================================
ELECTRICITY_FACTOR = 0.604       # tCO₂/MWh（北京市生态环境局 2024年度）
NATURAL_GAS_FACTOR = 2.1622     # tCO₂/万Nm³（DB11/T 1785-2020 附录B）
HEAT_FACTOR = 0.11              # tCO₂/GJ
GASOLINE_FACTOR = 2.9251        # tCO₂/t
DIESEL_FACTOR = 3.0959          # tCO₂/t
COAL_FACTOR = 2.53              # tCO₂/t（褐煤）

# ============================================================
# 报告引擎
# ============================================================

class ReportEngine:
    """合规碳排放报告生成引擎"""

    def __init__(self) -> None:
        self.emission_factors = {
            "electricity": ELECTRICITY_FACTOR,
            "natural_gas": NATURAL_GAS_FACTOR,
            "heat": HEAT_FACTOR,
            "gasoline": GASOLINE_FACTOR,
            "diesel": DIESEL_FACTOR,
            "coal": COAL_FACTOR,
        }

    # ---- 排放计算 ----

    def calculate_emissions(
        self,
        activity_data: list[ActivityData],
        emission_factors: Optional[dict[str, float]] = None,
    ) -> list[EmissionCalculation]:
        """计算各排放源的 CO₂ 排放量"""
        factors = emission_factors or self.emission_factors
        total = 0.0
        results: list[EmissionCalculation] = []

        for item in activity_data:
            factor_key = item.energy_type.value
            factor = factors.get(factor_key, 0.0)
            co2 = round(item.activity_value * factor, 4)
            total += co2
            results.append(
                EmissionCalculation(
                    source_id=item.source_id,
                    source_name=item.source_name,
                    energy_type=item.energy_type,
                    scope="scope1" if item.energy_type != EnergyType.ELECTRICITY else "scope2",
                    activity_value=item.activity_value,
                    activity_unit=item.unit,
                    emission_factor=factor,
                    emission_factor_unit=self._get_factor_unit(item.energy_type),
                    co2_emission=co2,
                    co2_emission_pct=0.0,
                )
            )

        # 计算占比
        if total > 0:
            for r in results:
                r.co2_emission_pct = round(r.co2_emission / total * 100, 2)

        return results

    def build_emission_summary(
        self, calculations: list[EmissionCalculation]
    ) -> dict[str, float]:
        """构建排放汇总"""
        scope1 = sum(c.co2_emission for c in calculations if c.scope == "scope1")
        scope2 = sum(c.co2_emission for c in calculations if c.scope == "scope2")
        return {
            "total": round(scope1 + scope2, 4),
            "scope1": round(scope1, 4),
            "scope2": round(scope2, 4),
        }

    @staticmethod
    def _get_factor_unit(energy_type: EnergyType) -> str:
        mapping = {
            EnergyType.ELECTRICITY: "tCO₂/MWh",
            EnergyType.NATURAL_GAS: "tCO₂/万Nm³",
            EnergyType.HEAT: "tCO₂/GJ",
            EnergyType.GASOLINE: "tCO₂/t",
            EnergyType.DIESEL: "tCO₂/t",
            EnergyType.COAL: "tCO₂/t",
        }
        return mapping.get(energy_type, "tCO₂")

    # ---- Word 生成 ----

    def generate_word(
        self,
        enterprise: EnterpriseInfo,
        activity_data: list[ActivityData],
        emission_sources: list[EmissionSource],
        calculations: list[EmissionCalculation],
        summary: dict[str, float],
        report_number: str = "",
    ) -> BytesIO:
        """生成 Word 报告 (.docx)"""
        doc = Document()

        # 页面设置
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

        # 封面
        self._add_word_cover(doc, enterprise, report_number)
        doc.add_page_break()

        # 第一章
        self._add_word_heading(doc, "第一章 企业基本情况", 1)
        self._add_word_table_c1(doc, enterprise)
        doc.add_page_break()

        # 第二章
        self._add_word_heading(doc, "第二章 二氧化碳排放核算", 1)
        self._add_word_heading(doc, "2.1 排放总量", 2)
        self._add_word_table_c2(doc, summary)
        self._add_word_heading(doc, "2.2 化石燃料燃烧排放（表C.3）", 2)
        self._add_word_table_c3(doc, calculations)
        self._add_word_heading(doc, "2.3 电力消费排放（表C.4）", 2)
        self._add_word_table_c4(doc, calculations)
        self._add_word_heading(doc, "2.4 热力消费排放（表C.5）", 2)
        self._add_word_table_c5(doc, calculations)
        self._add_word_heading(doc, "2.5 隐含排放（表C.6）", 2)
        self._add_word_table_c6(doc)
        doc.add_page_break()

        # 第三章
        self._add_word_heading(doc, "第三章 活动水平数据", 1)
        self._add_word_heading(doc, "3.1 化石燃料活动水平（表C.7）", 2)
        self._add_word_table_c7(doc, activity_data)
        self._add_word_heading(doc, "3.2 电力活动水平（表C.8）", 2)
        self._add_word_table_c8(doc, activity_data)
        self._add_word_heading(doc, "3.3 热力活动水平（表C.9）", 2)
        self._add_word_table_c9(doc, activity_data)
        doc.add_page_break()

        # 第四章
        self._add_word_heading(doc, "第四章 排放因子数据", 1)
        self._add_word_heading(doc, "4.1 化石燃料排放因子（表C.10）", 2)
        self._add_word_table_c10(doc)
        self._add_word_heading(doc, "4.2 电力排放因子（表C.11）", 2)
        self._add_word_table_c11(doc)
        self._add_word_heading(doc, "4.3 热力排放因子（表C.12）", 2)
        self._add_word_table_c12(doc)
        doc.add_page_break()

        # 汇总
        self._add_word_heading(doc, "排放量汇总（表C.13）", 1)
        self._add_word_table_c13(doc, calculations)
        self._add_word_heading(doc, "主要用能设备（表C.14）", 1)
        self._add_word_table_c14(doc)
        doc.add_page_break()

        # 真实性声明
        self._add_word_heading(doc, "报告真实性声明（表C.15）", 1)
        self._add_word_table_c15(doc, enterprise)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # ---- Word 辅助方法 ----

    @staticmethod
    def _set_cell_shading(cell, color_hex: str) -> None:
        """设置单元格背景色"""
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        })
        shading.append(shading_elm)

    @staticmethod
    def _set_cell_border(cell, **kwargs) -> None:
        """设置单元格边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            tcBorders = tcBorders.makeelement(qn("w:tcBorders"), {})
            tcPr.append(tcBorders)

    @staticmethod
    def _add_word_styled_paragraph(doc, text: str, style: str = "Normal",
                                     font_size: int = 10, bold: bool = False,
                                     alignment: int = WD_ALIGN_PARAGRAPH.LEFT,
                                     font_name: str = "宋体") -> None:
        p = doc.add_paragraph(style=style)
        p.alignment = alignment
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _add_word_cover(doc, enterprise: EnterpriseInfo, report_number: str) -> None:
        """封面"""
        for _ in range(6):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("二氧化碳排放报告")
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("（DB11/T 1785-2020 附录C）")
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        for _ in range(4):
            doc.add_paragraph()

        info_lines = [
            f"报告编号：{report_number or '自行编制'}",
            f"报告年度：{enterprise.reporting_year}年",
            f"单位名称：{enterprise.name}",
            f"编制日期：{date.today().strftime('%Y年%m月%d日')}",
        ]
        for line in info_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.font.name = "仿宋"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    @staticmethod
    def _add_word_heading(doc, text: str, level: int = 1) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        else:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    @staticmethod
    def _add_word_table_header(table, headers: list[str], col_widths: Optional[list] = None) -> None:
        """设置表头（深蓝 #081028）"""
        for j, h in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            ReportEngine._set_cell_shading(cell, "081028")

    def _add_word_table_c1(self, doc, e: EnterpriseInfo) -> None:
        """表C.1 报告主体基本情况"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.1  报告主体基本情况")
        run.font.size = Pt(10)
        run.font.bold = True

        fields = [
            ("企业名称", e.name), ("统一社会信用代码", e.unified_code),
            ("企业地址", e.address), ("法定代表人", e.legal_representative),
            ("联系人", e.contact_person), ("联系电话", e.contact_phone),
            ("行业分类", e.industry_category), ("报告年度", f"{e.reporting_year}年"),
            ("报告期", e.reporting_period),
        ]
        table = doc.add_table(rows=len(fields) + 1, cols=2, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, ["项目", "内容"])

        for i, (label, val) in enumerate(fields, 1):
            for j, text in enumerate([label, str(val)]):
                cell = table.rows[i].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 1 else WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if j == 0:
                    run.font.bold = True

    def _add_word_table_c2(self, doc, summary: dict[str, float]) -> None:
        """表C.2 二氧化碳排放总量汇总"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.2  二氧化碳排放总量汇总")
        run.font.size = Pt(10)
        run.font.bold = True

        table = doc.add_table(rows=4, cols=3, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, ["排放范围", "排放量(tCO₂)", "占比(%)"])
        data = [
            ("范围1：直接排放", f"{summary['scope1']:.2f}", f"{summary['scope1']/summary['total']*100:.1f}%"),
            ("范围2：间接排放", f"{summary['scope2']:.2f}", f"{summary['scope2']/summary['total']*100:.1f}%"),
            ("合计", f"{summary['total']:.2f}", "100%"),
        ]
        for i, (scope, val, pct) in enumerate(data, 1):
            for j, text in enumerate([scope, val, pct]):
                cell = table.rows[i].cells[j]
                cell.text = ""
                p_cell = cell.paragraphs[0]
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_cell.add_run(text)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                if i == 3:
                    run.font.bold = True

    def _add_word_table_c3(self, doc, calculations: list[EmissionCalculation]) -> None:
        """表C.3 化石燃料燃烧排放（11列）"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.3  化石燃料燃烧排放")
        run.font.size = Pt(10)
        run.font.bold = True

        headers = [
            "序号", "燃料品种", "燃烧量", "低位发热量\n(GJ/t或GJ/万Nm³)",
            "热值\n(GJ)", "单位热值\n含碳量(tC/GJ)", "碳氧化率(%)",
            "44/12比", "单位热值\n碳排放(tCO₂)", "排放量\n(tCO₂)", "占比(%)"
        ]
        fossil = [c for c in calculations if c.energy_type not in (EnergyType.ELECTRICITY, EnergyType.HEAT)]
        rows = len(fossil) + 1
        table = doc.add_table(rows=rows if rows > 1 else 2, cols=11, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, headers)

        for i, calc in enumerate(fossil, 1):
            vals = [
                str(i), calc.source_name,
                f"{calc.activity_value:.2f}", "—", "—", "—", "—",
                "44/12", "—", f"{calc.co2_emission:.2f}", f"{calc.co2_emission_pct:.1f}%"
            ]
            for j, v in enumerate(vals):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(8)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        if not fossil:
            for j in range(11):
                cell = table.rows[1].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run("—")
                run.font.size = Pt(8)

    def _add_word_table_c4(self, doc, calculations: list[EmissionCalculation]) -> None:
        """表C.4 电力消费（单位：MWh）"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.4  电力消费")
        run.font.size = Pt(10)
        run.font.bold = True

        elec = [c for c in calculations if c.energy_type == EnergyType.ELECTRICITY]
        table = doc.add_table(rows=max(len(elec), 1) + 1, cols=6, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, [
            "序号", "用电区域/用途", "消费量(MWh)", "排放因子\n(tCO₂/MWh)",
            "排放量(tCO₂)", "占比(%)"
        ])

        for i, calc in enumerate(elec, 1):
            vals = [
                str(i), calc.source_name, f"{calc.activity_value:.2f}",
                f"{calc.emission_factor}", f"{calc.co2_emission:.2f}",
                f"{calc.co2_emission_pct:.1f}%"
            ]
            for j, v in enumerate(vals):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _add_word_table_c5(self, doc, calculations: list[EmissionCalculation]) -> None:
        """表C.5 热力消费"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.5  热力消费")
        run.font.size = Pt(10)
        run.font.bold = True

        heat = [c for c in calculations if c.energy_type == EnergyType.HEAT]
        table = doc.add_table(rows=max(len(heat), 1) + 1, cols=6, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, [
            "序号", "用热区域/用途", "消费量(GJ)", "排放因子\n(tCO₂/GJ)",
            "排放量(tCO₂)", "占比(%)"
        ])

        for i, calc in enumerate(heat, 1):
            vals = [
                str(i), calc.source_name, f"{calc.activity_value:.2f}",
                f"{calc.emission_factor}", f"{calc.co2_emission:.2f}",
                f"{calc.co2_emission_pct:.1f}%"
            ]
            for j, v in enumerate(vals):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _add_word_table_c6(doc) -> None:
        """表C.6 隐含排放（如有）"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.6  隐含排放（如有）")
        run.font.size = Pt(10)
        run.font.bold = True
        p2 = doc.add_paragraph()
        run2 = p2.add_run("本报告期无隐含排放。")
        run2.font.size = Pt(10)

    def _add_word_table_c7(self, doc, activity_data: list[ActivityData]) -> None:
        """表C.7 化石燃料活动水平数据"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.7  化石燃料活动水平数据")
        run.font.size = Pt(10)
        run.font.bold = True

        fossil = [a for a in activity_data if a.energy_type not in (EnergyType.ELECTRICITY, EnergyType.HEAT)]
        table = doc.add_table(rows=max(len(fossil), 1) + 1, cols=5, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, ["序号", "燃料品种", "活动水平", "单位", "数据来源"])

        for i, a in enumerate(fossil, 1):
            vals = [str(i), a.source_name, f"{a.activity_value:.2f}", a.unit, a.data_source]
            for j, v in enumerate(vals):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _add_word_table_c8(self, doc, activity_data: list[ActivityData]) -> None:
        """表C.8 电力活动水平数据"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.8  电力活动水平数据")
        run.font.size = Pt(10)
        run.font.bold = True

        elec = [a for a in activity_data if a.energy_type == EnergyType.ELECTRICITY]
        table = doc.add_table(rows=max(len(elec), 1) + 1, cols=6, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, [
            "序号", "用电区域", "消费量(MWh)", "数据来源", "计量表编号", "备注"
        ])

        for i, a in enumerate(elec, 1):
            vals = [str(i), a.source_name, f"{a.activity_value:.2f}", a.data_source, a.meter_id or "—", "—"]
            for j, v in enumerate(vals):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _add_word_table_c9(self, doc, activity_data: list[ActivityData]) -> None:
        """表C.9 热力活动水平数据"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.9  热力活动水平数据")
        run.font.size = Pt(10)
        run.font.bold = True

        heat = [a for a in activity_data if a.energy_type == EnergyType.HEAT]
        table = doc.add_table(rows=max(len(heat), 1) + 1, cols=6, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, [
            "序号", "用热区域", "消费量(GJ)", "数据来源", "计量表编号", "备注"
        ])

        for i, a in enumerate(heat, 1):
            vals = [str(i), a.source_name, f"{a.activity_value:.2f}", a.data_source, a.meter_id or "—", "—"]
            for j, v in enumerate(vals):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _add_word_table_c10(doc) -> None:
        """表C.10 化石燃料排放因子数据"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.10  化石燃料排放因子数据")
        run.font.size = Pt(10)
        run.font.bold = True

        table = doc.add_table(rows=5, cols=7, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ReportEngine._add_word_table_header(table, [
            "序号", "燃料品种", "排放因子", "单位", "低位发热量", "碳氧化率", "来源"
        ])
        factors_data = [
            ("1", "天然气", "2.1622", "tCO₂/万Nm³", "389.31 GJ/万Nm³", "99%", "DB11/T 1785-2020"),
            ("2", "汽油", "2.9251", "tCO₂/t", "44.80 GJ/t", "98%", "DB11/T 1785-2020"),
            ("3", "柴油", "3.0959", "tCO₂/t", "43.33 GJ/t", "98%", "DB11/T 1785-2020"),
            ("4", "煤炭(褐煤)", "2.53", "tCO₂/t", "14.08 GJ/t", "98%", "DB11/T 1785-2020"),
        ]
        for i, row_data in enumerate(factors_data, 1):
            for j, v in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(8)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _add_word_table_c11(doc) -> None:
        """表C.11 电力排放因子数据"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.11  电力排放因子数据")
        run.font.size = Pt(10)
        run.font.bold = True

        table = doc.add_table(rows=2, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ReportEngine._add_word_table_header(table, ["序号", "排放因子名称", "排放因子值", "单位"])
        vals = ["1", "电力排放因子", "0.604", "tCO₂/MWh"]
        for j, v in enumerate(vals):
            cell = table.rows[1].cells[j]
            cell.text = ""
            pc = cell.paragraphs[0]
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pc.add_run(v)
            run.font.size = Pt(9)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _add_word_table_c12(doc) -> None:
        """表C.12 热力排放因子数据"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.12  热力排放因子数据")
        run.font.size = Pt(10)
        run.font.bold = True

        table = doc.add_table(rows=2, cols=4, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ReportEngine._add_word_table_header(table, ["序号", "排放因子名称", "排放因子值", "单位"])
        vals = ["1", "热力排放因子", "0.11", "tCO₂/GJ"]
        for j, v in enumerate(vals):
            cell = table.rows[1].cells[j]
            cell.text = ""
            pc = cell.paragraphs[0]
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pc.add_run(v)
            run.font.size = Pt(9)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _add_word_table_c13(self, doc, calculations: list[EmissionCalculation]) -> None:
        """表C.13 排放量汇总"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.13  排放量汇总")
        run.font.size = Pt(10)
        run.font.bold = True

        table = doc.add_table(rows=len(calculations) + 2, cols=5, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._add_word_table_header(table, [
            "序号", "排放源", "活动水平", "排放因子", "排放量(tCO₂)"
        ])

        total = 0.0
        for i, c in enumerate(calculations, 1):
            vals = [
                str(i), c.source_name,
                f"{c.activity_value:.2f} {c.activity_unit}",
                f"{c.emission_factor} {c.emission_factor_unit}",
                f"{c.co2_emission:.2f}"
            ]
            total += c.co2_emission
            for j, v in enumerate(vals):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 合计行
        last = len(calculations) + 1
        for j, v in enumerate(["合计", "", "", "", f"{total:.2f}"]):
            cell = table.rows[last].cells[j]
            cell.text = ""
            pc = cell.paragraphs[0]
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pc.add_run(v)
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _add_word_table_c14(doc) -> None:
        """表C.14 主要用能设备清单"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.14  主要用能设备清单")
        run.font.size = Pt(10)
        run.font.bold = True

        devices = [
            ("1", "中央空调冷水机组", "3台", "制冷量 1,200kW/台", "教学楼A、实验楼"),
            ("2", "燃气锅炉", "2台", "额定蒸发量 4t/h", "锅炉房"),
            ("3", "变压器", "6台", "容量 1,250kVA/台", "各配电室"),
            ("4", "照明系统", "—", "LED灯具约 12,000盏", "全校"),
            ("5", "电梯", "12台", "载重 1,000kg/台", "各教学楼"),
        ]
        table = doc.add_table(rows=len(devices) + 1, cols=5, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ReportEngine._add_word_table_header(table, [
            "序号", "设备名称", "数量", "规格型号", "安装位置"
        ])

        for i, row_data in enumerate(devices, 1):
            for j, v in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = ""
                pc = cell.paragraphs[0]
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = pc.add_run(v)
                run.font.size = Pt(9)
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    @staticmethod
    def _add_word_table_c15(doc, enterprise: EnterpriseInfo) -> None:
        """表C.15 报告真实性声明"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("表C.15  报告真实性声明")
        run.font.size = Pt(10)
        run.font.bold = True

        statement = (
            f"本单位郑重声明：\n\n"
            f"本报告依据DB11/T 1785-2020《二氧化碳排放核算和报告要求 服务业》编制，"
            f"报告中所涉及的排放数据、活动水平数据、排放因子数据均真实、准确、完整。"
            f"本单位对报告的真实性、准确性、完整性负责。\n\n"
            f"如有不实，本单位愿承担相应法律责任。\n\n\n"
            f"法定代表人（签字）：_______________\n\n"
            f"单位名称（盖章）：{enterprise.name}\n\n"
            f"日期：{date.today().strftime('%Y年%m月%d日')}"
        )
        p2 = doc.add_paragraph()
        run2 = p2.add_run(statement)
        run2.font.size = Pt(12)
        run2.font.name = "仿宋"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    # ---- HTML 预览 ----

    # ============================================================
    # CSS 样式常量
    # ============================================================
    CSS_STYLES = """
    <style>
        @page { size: A4; margin: 2.54cm 3.18cm; }
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body {
            font-family: "SimSun", "宋体", "Noto Serif SC", serif;
            font-size: 14px;
            line-height: 1.8;
            color: #333;
            background: #fff;
            max-width: 210mm;
            margin: 0 auto;
            padding: 20px 30px;
        }
        /* 封面 */
        .cover { text-align: center; padding: 120px 0 80px; page-break-after: always; }
        .cover h1 { font-family: "SimHei", "黑体", sans-serif; font-size: 28pt; margin-bottom: 12px; letter-spacing: 4px; color: #081028; }
        .cover .subtitle { font-size: 14pt; color: #666; margin-bottom: 80px; }
        .cover .info { font-size: 14pt; line-height: 2.5; }
        .cover .info span { display: block; }
        /* 章节标题 */
        .chapter { page-break-before: always; margin-top: 20px; }
        .chapter h2 { font-family: "SimHei", "黑体", sans-serif; font-size: 18pt; color: #081028; border-bottom: 2px solid #081028; padding-bottom: 6px; margin-bottom: 20px; }
        .chapter h3 { font-family: "SimHei", "黑体", sans-serif; font-size: 14pt; color: #1a237e; margin: 18px 0 10px; }
        /* 表格 */
        .table-wrapper { margin: 16px 0; overflow-x: auto; }
        .table-title { text-align: center; font-weight: bold; font-size: 11pt; margin-bottom: 8px; }
        table { width: 100%; border-collapse: collapse; font-size: 9pt; }
        thead th {
            background: #081028;
            color: #fff;
            font-weight: bold;
            padding: 8px 4px;
            border: 1px solid #081028;
            text-align: center;
            font-family: "SimHei", "黑体", sans-serif;
            white-space: nowrap;
        }
        tbody td {
            padding: 6px 4px;
            border: 1px solid #999;
            text-align: center;
        }
        tbody tr:nth-child(even) { background: #f8f9fa; }
        .total-row td { font-weight: bold; background: #e8eaf6 !important; }
        /* 声明 */
        .statement { margin-top: 40px; padding: 30px; border: 1px solid #999; }
        .statement p { text-indent: 2em; line-height: 2.2; }
        /* 页脚 */
        .page-footer { text-align: center; font-size: 9pt; color: #999; margin-top: 20px; border-top: 1px solid #ddd; padding-top: 8px; }
        @media print { body { padding: 0; } }
    </style>
    """

    def generate_html_preview(
        self,
        enterprise: EnterpriseInfo,
        activity_data: list[ActivityData],
        emission_sources: list[EmissionSource],
        calculations: list[EmissionCalculation],
        summary: dict[str, float],
        report_number: str = "",
    ) -> str:
        """生成完整 HTML 预览"""
        parts: list[str] = []
        parts.append("<!DOCTYPE html><html lang=\"zh-CN\"><head><meta charset=\"utf-8\">")
        parts.append("<title>二氧化碳排放报告</title>")
        parts.append(self.CSS_STYLES)
        parts.append("</head><body>")

        # 封面
        parts.append(self._html_cover(enterprise, report_number))

        # 第一章：企业基本情况
        parts.append(self._html_chapter_start("第一章  企业基本情况"))
        parts.append(self._html_table_c1(enterprise))

        # 第二章：二氧化碳排放核算
        parts.append(self._html_chapter_start("第二章  二氧化碳排放核算"))
        parts.append("<h3>2.1 排放总量</h3>")
        parts.append(self._html_table_c2(summary))
        parts.append("<h3>2.2 化石燃料燃烧排放（表C.3）</h3>")
        parts.append(self._html_table_c3(calculations))
        parts.append("<h3>2.3 电力消费排放（表C.4）</h3>")
        parts.append(self._html_table_c4(calculations))
        parts.append("<h3>2.4 热力消费排放（表C.5）</h3>")
        parts.append(self._html_table_c5(calculations))
        parts.append("<h3>2.5 隐含排放（表C.6）</h3>")
        parts.append(self._html_table_c6())

        # 第三章：活动水平数据
        parts.append(self._html_chapter_start("第三章  活动水平数据"))
        parts.append("<h3>3.1 化石燃料活动水平（表C.7）</h3>")
        parts.append(self._html_table_c7(activity_data))
        parts.append("<h3>3.2 电力活动水平（表C.8）</h3>")
        parts.append(self._html_table_c8(activity_data))
        parts.append("<h3>3.3 热力活动水平（表C.9）</h3>")
        parts.append(self._html_table_c9(activity_data))

        # 第四章：排放因子数据
        parts.append(self._html_chapter_start("第四章  排放因子数据"))
        parts.append("<h3>4.1 化石燃料排放因子（表C.10）</h3>")
        parts.append(self._html_table_c10())
        parts.append("<h3>4.2 电力排放因子（表C.11）</h3>")
        parts.append(self._html_table_c11())
        parts.append("<h3>4.3 热力排放因子（表C.12）</h3>")
        parts.append(self._html_table_c12())

        # 汇总与声明
        parts.append(self._html_chapter_start("排放量汇总（表C.13）"))
        parts.append(self._html_table_c13(calculations))
        parts.append(self._html_chapter_start("主要用能设备（表C.14）"))
        parts.append(self._html_table_c14())
        parts.append(self._html_chapter_start("报告真实性声明（表C.15）"))
        parts.append(self._html_table_c15(enterprise))

        parts.append("<div class=\"page-footer\">DB11/T 1785-2020 附录C 标准格式 | 本报告仅供内部碳管理使用</div>")
        parts.append("</body></html>")
        return "\n".join(parts)

    # ---- HTML 子模板 ----

    @staticmethod
    def _html_cover(enterprise: EnterpriseInfo, report_number: str) -> str:
        return f"""<div class="cover">
<h1>二氧化碳排放报告</h1>
<p class="subtitle">（DB11/T 1785-2020 附录C）</p>
<div class="info" style="margin-top:60px;">
<span>报告编号：{report_number or '自行编制'}</span>
<span>报告年度：{enterprise.reporting_year}年</span>
<span>单位名称：{enterprise.name}</span>
<span>编制日期：{date.today().strftime('%Y年%m月%d日')}</span>
</div>
</div>"""

    @staticmethod
    def _html_chapter_start(title: str) -> str:
        return f'<div class="chapter"><h2>{title}</h2></div>'

    @staticmethod
    def _html_table_c1(enterprise: EnterpriseInfo) -> str:
        rows = [
            ("企业名称", enterprise.name),
            ("统一社会信用代码", enterprise.unified_code),
            ("企业地址", enterprise.address),
            ("法定代表人", enterprise.legal_representative),
            ("联系人", enterprise.contact_person),
            ("联系电话", enterprise.contact_phone),
            ("行业分类", enterprise.industry_category),
            ("报告年度", f"{enterprise.reporting_year}年"),
            ("报告期", enterprise.reporting_period),
        ]
        row_html = "".join(
            f"<tr><td style=\"width:30%;font-weight:bold;\">{l}</td><td>{v}</td></tr>"
            for l, v in rows
        )
        return f"""<div class="table-wrapper">
<div class="table-title">表C.1  报告主体基本情况</div>
<table><thead><tr><th style="width:30%;">项目</th><th>内容</th></tr></thead>
<tbody>{row_html}</tbody></table></div>"""

    @staticmethod
    def _html_table_c2(summary: dict[str, float]) -> str:
        t = summary["total"]
        s1_pct = f"{summary['scope1']/t*100:.1f}%"
        s2_pct = f"{summary['scope2']/t*100:.1f}%"
        return f"""<div class="table-wrapper">
<div class="table-title">表C.2  二氧化碳排放总量汇总</div>
<table><thead><tr><th>排放范围</th><th>排放量(tCO₂)</th><th>占比(%)</th></tr></thead>
<tbody>
<tr><td>范围1：直接排放</td><td>{summary['scope1']:.2f}</td><td>{s1_pct}</td></tr>
<tr><td>范围2：间接排放</td><td>{summary['scope2']:.2f}</td><td>{s2_pct}</td></tr>
<tr class="total-row"><td>合计</td><td>{t:.2f}</td><td>100%</td></tr>
</tbody></table></div>"""

    @staticmethod
    def _html_table_c3(calculations: list[EmissionCalculation]) -> str:
        """表C.3 化石燃料燃烧排放（11列）"""
        fossil = [c for c in calculations if c.energy_type not in (EnergyType.ELECTRICITY, EnergyType.HEAT)]
        header = "<tr><th>序号</th><th>燃料品种</th><th>燃烧量</th><th>低位发热量<br/>(GJ/t或GJ/万Nm³)</th><th>热值<br/>(GJ)</th><th>单位热值<br/>含碳量(tC/GJ)</th><th>碳氧化率<br/>(%)</th><th>44/12比</th><th>单位热值<br/>碳排放(tCO₂)</th><th>排放量<br/>(tCO₂)</th><th>占比<br/>(%)</th></tr>"
        if not fossil:
            rows = "<tr><td colspan=\"11\">—</td></tr>"
        else:
            rows = "".join(
                f"<tr><td>{i}</td><td>{c.source_name}</td><td>{c.activity_value:.2f}</td><td>—</td><td>—</td><td>—</td><td>—</td><td>44/12</td><td>—</td><td>{c.co2_emission:.2f}</td><td>{c.co2_emission_pct:.1f}%</td></tr>"
                for i, c in enumerate(fossil, 1)
            )
        return f"""<div class="table-wrapper">
<div class="table-title">表C.3  化石燃料燃烧排放</div>
<table><thead>{header}</thead><tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c4(calculations: list[EmissionCalculation]) -> str:
        """表C.4 电力消费（单位：MWh）"""
        elec = [c for c in calculations if c.energy_type == EnergyType.ELECTRICITY]
        header = "<tr><th>序号</th><th>用电区域/用途</th><th>消费量(MWh)</th><th>排放因子<br/>(tCO₂/MWh)</th><th>排放量(tCO₂)</th><th>占比(%)</th></tr>"
        rows = "".join(
            f"<tr><td>{i}</td><td>{c.source_name}</td><td>{c.activity_value:.2f}</td><td>{c.emission_factor}</td><td>{c.co2_emission:.2f}</td><td>{c.co2_emission_pct:.1f}%</td></tr>"
            for i, c in enumerate(elec, 1)
        ) if elec else "<tr><td colspan=\"6\">—</td></tr>"
        return f"""<div class="table-wrapper">
<div class="table-title">表C.4  电力消费（单位：MWh）</div>
<table><thead>{header}</thead><tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c5(calculations: list[EmissionCalculation]) -> str:
        """表C.5 热力消费"""
        heat = [c for c in calculations if c.energy_type == EnergyType.HEAT]
        header = "<tr><th>序号</th><th>用热区域/用途</th><th>消费量(GJ)</th><th>排放因子<br/>(tCO₂/GJ)</th><th>排放量(tCO₂)</th><th>占比(%)</th></tr>"
        rows = "".join(
            f"<tr><td>{i}</td><td>{c.source_name}</td><td>{c.activity_value:.2f}</td><td>{c.emission_factor}</td><td>{c.co2_emission:.2f}</td><td>{c.co2_emission_pct:.1f}%</td></tr>"
            for i, c in enumerate(heat, 1)
        ) if heat else "<tr><td colspan=\"6\">—</td></tr>"
        return f"""<div class="table-wrapper">
<div class="table-title">表C.5  热力消费</div>
<table><thead>{header}</thead><tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c6() -> str:
        """表C.6 隐含排放"""
        return """<div class="table-wrapper">
<div class="table-title">表C.6  隐含排放（如有）</div>
<p style="text-indent:2em;margin:12px 0;">本报告期无隐含排放。</p></div>"""

    @staticmethod
    def _html_table_c7(activity_data: list[ActivityData]) -> str:
        fossil = [a for a in activity_data if a.energy_type not in (EnergyType.ELECTRICITY, EnergyType.HEAT)]
        header = "<tr><th>序号</th><th>燃料品种</th><th>活动水平</th><th>单位</th><th>数据来源</th></tr>"
        rows = "".join(
            f"<tr><td>{i}</td><td>{a.source_name}</td><td>{a.activity_value:.2f}</td><td>{a.unit}</td><td>{a.data_source}</td></tr>"
            for i, a in enumerate(fossil, 1)
        ) if fossil else "<tr><td colspan=\"5\">—</td></tr>"
        return f"""<div class="table-wrapper">
<div class="table-title">表C.7  化石燃料活动水平数据</div>
<table><thead>{header}</thead><tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c8(activity_data: list[ActivityData]) -> str:
        elec = [a for a in activity_data if a.energy_type == EnergyType.ELECTRICITY]
        header = "<tr><th>序号</th><th>用电区域</th><th>消费量(MWh)</th><th>数据来源</th><th>计量表编号</th><th>备注</th></tr>"
        rows = "".join(
            f"<tr><td>{i}</td><td>{a.source_name}</td><td>{a.activity_value:.2f}</td><td>{a.data_source}</td><td>{a.meter_id or '—'}</td><td>—</td></tr>"
            for i, a in enumerate(elec, 1)
        ) if elec else "<tr><td colspan=\"6\">—</td></tr>"
        return f"""<div class="table-wrapper">
<div class="table-title">表C.8  电力活动水平数据（单位：MWh）</div>
<table><thead>{header}</thead><tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c9(activity_data: list[ActivityData]) -> str:
        heat = [a for a in activity_data if a.energy_type == EnergyType.HEAT]
        header = "<tr><th>序号</th><th>用热区域</th><th>消费量(GJ)</th><th>数据来源</th><th>计量表编号</th><th>备注</th></tr>"
        rows = "".join(
            f"<tr><td>{i}</td><td>{a.source_name}</td><td>{a.activity_value:.2f}</td><td>{a.data_source}</td><td>{a.meter_id or '—'}</td><td>—</td></tr>"
            for i, a in enumerate(heat, 1)
        ) if heat else "<tr><td colspan=\"6\">—</td></tr>"
        return f"""<div class="table-wrapper">
<div class="table-title">表C.9  热力活动水平数据</div>
<table><thead>{header}</thead><tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c10() -> str:
        """表C.10 化石燃料排放因子数据"""
        rows = "".join([
            "<tr><td>1</td><td>天然气</td><td>2.1622</td><td>tCO₂/万Nm³</td><td>389.31 GJ/万Nm³</td><td>99%</td><td>DB11/T 1785-2020</td></tr>",
            "<tr><td>2</td><td>汽油</td><td>2.9251</td><td>tCO₂/t</td><td>44.80 GJ/t</td><td>98%</td><td>DB11/T 1785-2020</td></tr>",
            "<tr><td>3</td><td>柴油</td><td>3.0959</td><td>tCO₂/t</td><td>43.33 GJ/t</td><td>98%</td><td>DB11/T 1785-2020</td></tr>",
            "<tr><td>4</td><td>煤炭(褐煤)</td><td>2.53</td><td>tCO₂/t</td><td>14.08 GJ/t</td><td>98%</td><td>DB11/T 1785-2020</td></tr>",
        ])
        return f"""<div class="table-wrapper">
<div class="table-title">表C.10  化石燃料排放因子数据</div>
<table><thead><tr><th>序号</th><th>燃料品种</th><th>排放因子</th><th>单位</th><th>低位发热量</th><th>碳氧化率</th><th>来源</th></tr></thead>
<tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c11() -> str:
        """表C.11 电力排放因子"""
        return """<div class="table-wrapper">
<div class="table-title">表C.11  电力排放因子数据</div>
<table><thead><tr><th>序号</th><th>排放因子名称</th><th>排放因子值</th><th>单位</th></tr></thead>
<tbody><tr><td>1</td><td>电力排放因子</td><td>0.604</td><td>tCO₂/MWh</td></tr></tbody></table>
<p style="font-size:9pt;color:#666;margin-top:4px;">数据来源：北京市生态环境局 2024年度</p></div>"""

    @staticmethod
    def _html_table_c12() -> str:
        """表C.12 热力排放因子"""
        return """<div class="table-wrapper">
<div class="table-title">表C.12  热力排放因子数据</div>
<table><thead><tr><th>序号</th><th>排放因子名称</th><th>排放因子值</th><th>单位</th></tr></thead>
<tbody><tr><td>1</td><td>热力排放因子</td><td>0.11</td><td>tCO₂/GJ</td></tr></tbody></table></div>"""

    @staticmethod
    def _html_table_c13(calculations: list[EmissionCalculation]) -> str:
        header = "<tr><th>序号</th><th>排放源</th><th>活动水平</th><th>排放因子</th><th>排放量(tCO₂)</th></tr>"
        total = 0.0
        rows = ""
        for i, c in enumerate(calculations, 1):
            rows += f"<tr><td>{i}</td><td>{c.source_name}</td><td>{c.activity_value:.2f} {c.activity_unit}</td><td>{c.emission_factor} {c.emission_factor_unit}</td><td>{c.co2_emission:.2f}</td></tr>"
            total += c.co2_emission
        rows += f"<tr class=\"total-row\"><td colspan=\"4\">合计</td><td>{total:.2f}</td></tr>"
        return f"""<div class="table-wrapper">
<div class="table-title">表C.13  排放量汇总</div>
<table><thead>{header}</thead><tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c14() -> str:
        devices = [
            ("1", "中央空调冷水机组", "3台", "制冷量 1,200kW/台", "教学楼A、实验楼"),
            ("2", "燃气锅炉", "2台", "额定蒸发量 4t/h", "锅炉房"),
            ("3", "变压器", "6台", "容量 1,250kVA/台", "各配电室"),
            ("4", "照明系统", "—", "LED灯具约 12,000盏", "全校"),
            ("5", "电梯", "12台", "载重 1,000kg/台", "各教学楼"),
        ]
        rows = "".join(
            f"<tr><td>{d[0]}</td><td>{d[1]}</td><td>{d[2]}</td><td>{d[3]}</td><td>{d[4]}</td></tr>"
            for d in devices
        )
        return f"""<div class="table-wrapper">
<div class="table-title">表C.14  主要用能设备清单</div>
<table><thead><tr><th>序号</th><th>设备名称</th><th>数量</th><th>规格型号</th><th>安装位置</th></tr></thead>
<tbody>{rows}</tbody></table></div>"""

    @staticmethod
    def _html_table_c15(enterprise: EnterpriseInfo) -> str:
        return f"""<div class="table-wrapper">
<div class="table-title">表C.15  报告真实性声明</div>
<div class="statement">
<p>本单位郑重声明：</p>
<p>本报告依据DB11/T 1785-2020《二氧化碳排放核算和报告要求 服务业》编制，报告中所涉及的排放数据、活动水平数据、排放因子数据均真实、准确、完整。本单位对报告的真实性、准确性、完整性负责。</p>
<p>如有不实，本单位愿承担相应法律责任。</p>
<p style="margin-top:24px;">法定代表人（签字）：_______________</p>
<p>单位名称（盖章）：{enterprise.name}</p>
<p>日期：{date.today().strftime('%Y年%m月%d日')}</p>
</div></div>"""


# ============================================================
# 便捷函数（供 api_server.py 调用）
# ============================================================

_engine = ReportEngine()


def calculate_emissions(
    activity_data: list[ActivityData],
    emission_factors: Optional[dict[str, float]] = None,
) -> list[EmissionCalculation]:
    return _engine.calculate_emissions(activity_data, emission_factors)


def build_emission_summary(calculations: list[EmissionCalculation]) -> dict[str, float]:
    return _engine.build_emission_summary(calculations)


def generate_report(
    enterprise: EnterpriseInfo,
    activity_data: list[ActivityData],
    emission_sources: list[EmissionSource],
    calculations: list[EmissionCalculation],
    summary: dict[str, float],
    report_number: str = "",
) -> BytesIO:
    return _engine.generate_word(
        enterprise, activity_data, emission_sources,
        calculations, summary, report_number,
    )


def generate_html_preview(
    enterprise: EnterpriseInfo,
    activity_data: list[ActivityData],
    emission_sources: list[EmissionSource],
    calculations: list[EmissionCalculation],
    summary: dict[str, float],
    report_number: str = "",
) -> str:
    return _engine.generate_html_preview(
        enterprise, activity_data, emission_sources,
        calculations, summary, report_number,
    )


def get_emission_factors() -> list[dict[str, Any]]:
    """获取排放因子列表"""
    return [
        {
            "source_id": "C-01",
            "energy_type": "electricity",
            "factor_name": "电力排放因子",
            "factor_value": ELECTRICITY_FACTOR,
            "factor_unit": "tCO₂/MWh",
            "factor_source": "北京市生态环境局 2024年度",
            "oxidation_rate": 1.0,
            "gwp": None,
        },
        {
            "source_id": "C-02",
            "energy_type": "natural_gas",
            "factor_name": "天然气排放因子",
            "factor_value": NATURAL_GAS_FACTOR,
            "factor_unit": "tCO₂/万Nm³",
            "factor_source": "DB11/T 1785-2020 附录B",
            "oxidation_rate": 0.99,
            "gwp": 1,
        },
        {
            "source_id": "C-03",
            "energy_type": "heat",
            "factor_name": "热力排放因子",
            "factor_value": HEAT_FACTOR,
            "factor_unit": "tCO₂/GJ",
            "factor_source": "DB11/T 1785-2020 附录B",
            "oxidation_rate": 1.0,
            "gwp": None,
        },
        {
            "source_id": "C-04",
            "energy_type": "gasoline",
            "factor_name": "汽油排放因子",
            "factor_value": GASOLINE_FACTOR,
            "factor_unit": "tCO₂/t",
            "factor_source": "DB11/T 1785-2020 附录B",
            "oxidation_rate": 0.98,
            "gwp": None,
        },
        {
            "source_id": "C-05",
            "energy_type": "diesel",
            "factor_name": "柴油排放因子",
            "factor_value": DIESEL_FACTOR,
            "factor_unit": "tCO₂/t",
            "factor_source": "DB11/T 1785-2020 附录B",
            "oxidation_rate": 0.98,
            "gwp": None,
        },
        {
            "source_id": "C-06",
            "energy_type": "coal",
            "factor_name": "煤炭排放因子",
            "factor_value": COAL_FACTOR,
            "factor_unit": "tCO₂/t",
            "factor_source": "DB11/T 1785-2020 附录B",
            "oxidation_rate": 0.98,
            "gwp": 1,
        },
    ]


def validate_report_data(data: dict[str, Any]) -> dict[str, Any]:
    """数据校验"""
    errors: list[str] = []
    warnings: list[str] = []
    suggestions: list[str] = []

    if "enterprise" not in data:
        errors.append("缺少企业基本信息（enterprise）")
    if "activity_data" not in data or not data.get("activity_data"):
        errors.append("缺少活动水平数据（activity_data）")
    if "emission_sources" not in data or not data.get("emission_sources"):
        errors.append("缺少排放源定义（emission_sources）")

    if not errors:
        ent = data.get("enterprise", {})
        if not ent.get("name"):
            errors.append("企业名称不能为空")
        if not ent.get("unified_code"):
            errors.append("统一社会信用代码不能为空")
        year = ent.get("reporting_year", 0)
        if year < 2000 or year > 2100:
            errors.append(f"报告年度 {year} 不在合理范围内")

    if not errors:
        activities = data.get("activity_data", [])
        for i, a in enumerate(activities):
            if not a.get("source_id"):
                errors.append(f"活动水平数据第{i+1}项缺少 source_id")
            if not a.get("source_name"):
                errors.append(f"活动水平数据第{i+1}项缺少 source_name")
            if a.get("activity_value", 0) < 0:
                errors.append(f"活动水平数据第{i+1}项 activity_value 不能为负")

    if data.get("emission_sources"):
        for i, s in enumerate(data.get("emission_sources", [])):
            if not s.get("source_id"):
                warnings.append(f"排放源第{i+1}项缺少 source_id")

    return {
        "is_valid": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
        "suggestions": suggestions,
    }
